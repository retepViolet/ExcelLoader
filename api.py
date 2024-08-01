from fastapi import FastAPI, File, UploadFile, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from contextlib import asynccontextmanager
from prisma import Prisma
from typing import List
from formulas import ExcelModel
from docx import Document
import re, os, json, random, aiohttp, io, asyncio

path = 'C:\\Users\\fudan\\Desktop\\excel loader\\src\\api'
db = Prisma()
model_mem = {}     # 对计算模型的缓存
lock_model_mem = asyncio.Lock()     # 防止线程冲突

@asynccontextmanager
async def lifespan(_app: FastAPI):
    await db.connect()
    yield
    await db.disconnect()
app = FastAPI(lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 允许所有来源，生产环境中请更加谨慎地配置
    allow_credentials=True,
    allow_methods=["*"],  # 允许所有 HTTP 方法
    allow_headers=["*"],  # 允许所有 HTTP 头
)


def get_alpha(num):
    # Converts a column number to its corresponding Excel column letter
    result = ''
    while num > 0:
        num, remainder = divmod(num - 1, 26)
        result = chr(remainder + 65) + result
    return result

def get_num(cell):
    # Splits an Excel cell address into its column and row numbers
    column = ''.join(filter(str.isalpha, cell))
    row = ''.join(filter(str.isdigit, cell))
    num = 0
    for c in column:
        num = num * 26 + (ord(c.upper()) - ord('A') + 1)
    return num, int(row)

def get_id(fpath, sheet, cell: str):
    if ':' not in cell:
        return [f"'[{fpath}]{sheet.upper()}'!{cell.upper()}"]
    if ',' in cell:
        cells = cell.split(',')
        res = []
        for cell in cells: 
            res.extend(get_id(fpath, sheet, cell))
        return res
    
    start_cell, end_cell = cell.split(':')
    start_x, start_y = get_num(start_cell)
    end_x, end_y = get_num(end_cell)
    
    # Ensure start_x, start_y <= end_x, end_y
    if start_x > end_x:
        start_x, end_x = end_x, start_x
    if start_y > end_y:
        start_y, end_y = end_y, start_y
    
    res = []
    for y in range(start_y, end_y + 1):
        for x in range(start_x, end_x + 1):
            cell_id = f"'[{fpath}]{sheet.upper()}'!{get_alpha(x)}{y}"
            res.append(cell_id)
    
    return res

    

@app.post('/upload')
async def upload(fpath: str = Query()):
    try: model = ExcelModel().loads(fpath).finish()
    except: return f'Error: Fail to load "{fpath}".'
    name = os.path.basename(fpath)
    version = len(await db.file.find_many(where={'name': name}, order={'version': 'desc'})) + 1
    await db.file.create(data={'name': name, 'version': version, 'model': json.dumps(model.to_dict())})
    return {
        'file name': name,
        'version': version
    }


@app.delete('/delete')
async def delete(name: str = Query()):
    cnt = await db.file.delete_many(where={'name':name})
    return f'{cnt} number of excel files are deleted.'


async def load_model(file_name, version):
    if version < 0:
        return 'Error: Version should be a positive number.'
    
    files = await db.file.find_many(where={'name': file_name}, order={'version': 'desc'})
    if version == 0: version = len(files)
    if len(files) == 0: return f'Error: Please upload "{file_name}" first.'
    if len(files) < version: return f'Error: No such a version for "{file_name}".'
    file_id = files[version - 1].id
    if file_id in model_mem:
        model = model_mem[file_id]
        print('I have memory.')
    else:
        model = ExcelModel().from_dict(json.loads(files[version - 1].model))
        async with lock_model_mem:
            model_mem[file_id] = model
    return model, file_id, version


def get_input_output(file_name, input_cell, output_cell):
    try:
        input_cell = json.loads(input_cell)
        output_cell = json.loads(output_cell)
    except: 
        return 'Error: Cannot load input and output. They should be json string.'
    if not isinstance(input_cell, list) or not isinstance(output_cell, list):
        return 'Error: Input and output should be list.'
    
    input, output = {}, []
    for cell in input_cell:
        if 'sheet' not in cell or 'cell' not in cell or 'value' not in cell:
            return f'Error: Missing parameter as an input cell in {cell}.'
        try: ids = get_id(file_name, cell['sheet'], cell['cell'])
        except: return f'Error: Cell ID mistake in {cell}.'
        try: value = float(cell['value'])
        except: return f'Error: Value is not a number in {cell}.'
        for id in ids:
            input[id] = value
    for cell in output_cell:
        if 'sheet' not in cell or 'cell' not in cell:
            return f'Error: Missing parameter as an output cell in {cell}.'
        try: ids = get_id(file_name, cell['sheet'], cell['cell'])
        except: return f'Error: Cell ID mistake in {cell}.'
        output.extend(ids)
    return input, output


def get_result(model, input, output):
    solution, res = model.calculate(inputs = input, outputs = output), {}
    for cell in output:
        if cell in solution:
            res[cell] = solution[cell].value[0][0]
        else: res[cell] = None
    return res


def extract_identifiers(doc, identifier_pattern = r'\{\{.*?\}\}'):
    identifiers = []
    for para in doc.paragraphs:
        matches = re.findall(identifier_pattern, para.text)
        for match in matches:
            identifiers.append((para, match))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(identifier_pattern, para.text)
                    for match in matches:
                        identifiers.append((para, match))
    return identifiers
 

def write_docx(input_path, output_path, replacements):
    try: doc = Document(input_path)
    except: return f'Error: Cannot load file from "{input_path}".'
    identifiers = extract_identifiers(doc)
    while identifiers:
        para, cell_id = identifiers.pop(0)
        key = cell_id[2:-2]
        if key in replacements:
            para.text = para.text.replace(cell_id, str(round(replacements[key], 2)))
    try: doc.save(output_path)
    except: return f'Error: Cannot save file to "{output_path}".'


def get_result_json(ele, res, fpath):
    buf = []
    cell_ids = get_id(fpath, **ele)
    for cell_id in cell_ids:
        buf.append(res[cell_id])
    return buf


def get_json_output(output_json, res, file_name):
    if isinstance(output_json, (dict, list)):
        for key in (output_json if isinstance(output_json, dict) else range(len(output_json))):
            ele = output_json[key]
            if isinstance(ele, dict) and 'sheet' in ele and 'cell' in ele:
                value = get_result_json(ele, res, file_name)
                output_json[key] = value if len(value)>1 else value[0]
            else: get_json_output(ele, res, file_name)


@app.post('/calculate')
async def calculate(file_name: str = Query(), version: int = Query(0),
                    input_cell: str = Query(), output_cell: str = Query(),
                    output_excel: str = Query(''), output_docx: str = Query(''), 
                    output_json: str = Query(''))->dict:
    """
    input_cell: [
        { 
            'sheet': 'SHEET_NAME', 
            'cell': 'CELL or CELL_RANGE', 
            'value': 'INPUT_VALUE'
        }
    ]

    output_cell: [
        {
            'sheet': 'SHEET_NAME', 
            'cell': 'CELL or CELL_RANGE', 
        }
    ]

    output_excel: None or path to save the file.

    output_docx: {
        'input_path': 'Path to the template file.'
        'output_path': 'Where to save the new file.'
    }
    """

    buf = await load_model(file_name, version)
    if isinstance(buf, str):
        return buf
    model, file_id, version = buf

    buf = get_input_output(file_name, input_cell, output_cell)
    if isinstance(buf, str):
        return buf
    input, output = buf

    res = get_result(model, input, output)
    await db.history.create(data={
        'file_id': file_id,
        'intput': input_cell,
        'output': output_cell
    })

    if output_excel != '':
        try: model.write(dirpath = output_excel)
        except: return f'Error: Cannot write excel into "{output_excel}".'

    if output_docx != '':
        try: output_docx = json.loads(output_docx)
        except: return 'Error: Parameter output_docx should be in json format.'
        buf = write_docx(output_docx['input_path'], output_docx['output_path'], res)
        if isinstance(buf, str):
            return buf
        
    if output_json != '':
        try: output_json = json.loads(output_json)
        except: return 'Error: Parameter output_json should be in json format.'
        buf = get_json_output(output_json, res, file_name)
        print(output_json)
        return output_json
    
    return {
        'input': input,
        'output': res,
        'file name': file_name,
        'version': version
    }


@app.get('/history')
async def history(file_name: str = Query(), version: int = Query(0)):
    files = await db.file.find_many(where={'name': file_name}, order={'version': 'desc'}, include={'history': True})
    if version == 0: version = len(files)
    if len(files) == 0: return f'Error: Please upload "{file_name}" first.'
    if len(files) < version: return f'Error: No such a version for "{file_name}".'
    file = files[version - 1]
    return file.history


@app.get('/file_list')
async def file_list():
    files = await db.file.find_many()
    for file in files:
        file.model = None
    return files
